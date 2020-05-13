#!/usr/bin/python
# -*- coding: utf-8 -*-
import mysql.connector
from docx import Document


db = input("请输入数据库名称： ")
host = 'localhost'
user = 'root'
password = ''
port = 3306

conn = mysql.connector.connect(host=host, user=user, passwd=password, db=db, port=port)
cursor = conn.cursor()
cursor.execute(f"select table_name,table_comment from information_schema.tables where table_schema = '{db}' and table_comment != ''")
document = Document()

for (table_name, table_comment) in cursor.fetchall():
    document.add_heading(table_name, level=1)
    document.add_paragraph(table_comment)
    table = document.add_table(rows=1, cols=3, style='Table Grid')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = u'字段'
    hdr_cells[1].text = u'类型'
    hdr_cells[2].text = u'说明'
    cursor.execute(f"SHOW FULL FIELDS FROM {db}.%s" % table_name)
    for (field, type, _, _, _, _, _, _, comment) in cursor.fetchall():
        row_cells = table.add_row().cells
        row_cells[0].text = field
        row_cells[1].text = type
        row_cells[2].text = comment

document.save('db.docx')